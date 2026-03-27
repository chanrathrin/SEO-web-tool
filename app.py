from __future__ import annotations

import base64
import html
import io
import os
import re
import tempfile
import uuid
from dataclasses import dataclass
from typing import Any

from flask import Flask, jsonify, render_template, request, send_file
from PIL import Image

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.shared import RGBColor
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024


# -----------------------------
# Helpers
# -----------------------------
def clean_lines(text: str) -> list[str]:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\u00a0", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    raw_lines = [line.strip() for line in text.split("\n")]
    cleaned: list[str] = []
    last_blank = True

    for line in raw_lines:
        if not line:
            if not last_blank:
                cleaned.append("")
            last_blank = True
            continue
        cleaned.append(line)
        last_blank = False

    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    return cleaned


def trim_at_word_boundary(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if len(text) <= limit:
        return text
    cut = text[:limit].rstrip()
    if " " in cut:
        cut = cut.rsplit(" ", 1)[0]
    return cut.rstrip(" ,.-:;")


def strip_internal_seo_lines(lines: list[str]) -> list[str]:
    seo_prefixes = (
        "Focus Keyphrase:",
        "SEO Title:",
        "Meta Description:",
        "Slug (URL):",
        "Slug:",
        "Short Summary:",
    )
    cleaned = []
    for line in lines:
        s = line.strip()
        if any(s.startswith(prefix) for prefix in seo_prefixes):
            continue
        cleaned.append(line)

    while cleaned and cleaned[-1] == "":
        cleaned.pop()
    return cleaned


def guess_title(lines: list[str]) -> str:
    if not lines:
        return "Untitled Article"
    title = lines[0].strip()
    return title[:140].strip() if len(title) > 140 else title


def build_intro(lines: list[str]) -> str:
    content = [line for line in lines[1:] if line.strip()]
    if not content:
        return ""
    intro_parts = []
    for line in content:
        intro_parts.append(line)
        joined = " ".join(intro_parts).strip()
        if len(joined) >= 180 or line.endswith((".", "!", "?")):
            break
    return trim_at_word_boundary(" ".join(intro_parts).strip(), 240)


def split_body_into_sections(lines: list[str], num_sections: int = 3) -> list[str]:
    content_lines = [line.strip() for line in lines[1:] if line.strip()]
    if not content_lines:
        return []

    paragraphs = []
    current = []
    for line in lines[1:]:
        if not line.strip():
            if current:
                paragraphs.append(" ".join(current).strip())
                current = []
        else:
            current.append(line.strip())
    if current:
        paragraphs.append(" ".join(current).strip())

    blocks = [p for p in paragraphs if p] or content_lines
    if len(blocks) <= num_sections:
        return blocks[:num_sections]

    target = min(num_sections, len(blocks))
    base = len(blocks) // target
    extra = len(blocks) % target
    sections = []
    idx = 0
    for i in range(target):
        take = base + (1 if i < extra else 0)
        chunk = blocks[idx:idx + take]
        idx += take
        merged = "\n\n".join(chunk).strip()
        if merged:
            sections.append(merged)
    return sections[:num_sections]


def clean_heading_candidate(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip(' "\'“”‘’.,:;!?-')
    text = re.sub(r"^[^A-Za-z0-9]+", "", text)
    text = re.sub(r"[^A-Za-z0-9]+$", "", text)
    return text


def sentence_candidates_from_text(text: str) -> list[str]:
    text = text.replace("\n", " ")
    sentences = re.split(r'(?<=[.!?])\s+', text)
    out = []
    for s in sentences:
        s = clean_heading_candidate(s)
        if not s:
            continue
        wc = len(s.split())
        if 4 <= wc <= 12:
            out.append(s)
    return out


def phrase_candidates_from_text(text: str) -> list[str]:
    words = [w for w in clean_heading_candidate(text).split() if w]
    candidates = []
    for length in (5, 6, 7, 8):
        if len(words) >= length:
            candidates.append(" ".join(words[:length]))
    return candidates


def choose_heading_from_text(text: str, seen: set[str]) -> str:
    candidates = sentence_candidates_from_text(text)
    if not candidates:
        candidates = phrase_candidates_from_text(text)

    weak_starts = {
        "the", "a", "an", "this", "that", "these", "those",
        "here", "there", "it", "he", "she", "they", "we"
    }
    ranked = []
    for cand in candidates:
        cleaned = clean_heading_candidate(cand)
        if not cleaned:
            continue
        key = cleaned.lower()
        words = cleaned.split()
        if len(words) < 4:
            continue
        score = 0
        if 5 <= len(words) <= 10:
            score += 4
        elif len(words) <= 12:
            score += 2
        if not cleaned.endswith((":", ",", "-")):
            score += 1
        if words[0].lower() not in weak_starts:
            score += 2
        ranked.append((score, cleaned, key))

    ranked.sort(key=lambda x: (-x[0], abs(len(x[1]) - 52), len(x[1])))

    for _, cand, key in ranked:
        if key in seen:
            continue
        seen.add(key)
        return cand
    return ""


def build_nested_article_structure(sections: list[str]) -> list[dict[str, Any]]:
    structure = []
    seen: set[str] = set()
    cleaned_sections = [s.strip() for s in sections if s.strip()]
    if not cleaned_sections:
        return structure

    for idx, section in enumerate(cleaned_sections[:3], start=1):
        parts = [part.strip() for part in section.split("\n\n") if part.strip()]
        if not parts:
            continue

        search_pool = parts[:2] + [section]
        h2 = ""
        for candidate_source in search_pool:
            h2 = choose_heading_from_text(candidate_source, seen)
            if h2:
                break
        if not h2:
            fallback_words = clean_heading_candidate(parts[0]).split()
            h2 = " ".join(fallback_words[:8]).strip() or f"Section {idx}"

        subsections = []
        if len(parts) >= 3:
            lead_body = parts[0]
            mid_body = "\n\n".join(parts[1:-1]).strip() if len(parts) > 3 else parts[1]
            end_body = parts[-1]
            subsections.append({"h3": "", "h4": "", "body": lead_body})
            if mid_body and mid_body != lead_body and mid_body != end_body:
                subsections.append({"h3": "", "h4": "", "body": mid_body})
            if end_body and end_body != lead_body:
                subsections.append({"h3": "", "h4": "", "body": end_body})
        else:
            subsections.append({"h3": "", "h4": "", "body": "\n\n".join(parts).strip()})

        structure.append({"h2": h2, "subsections": subsections})

    return structure[:3]


def make_slug(title: str) -> str:
    slug = re.sub(r"[^a-z0-9\s-]", "", title.lower())
    slug = re.sub(r"\s+", "-", slug).strip("-")
    slug = re.sub(r"-{2,}", "-", slug)
    return slug


def make_focus_keyphrase(title: str) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", title).strip()
    words = cleaned.split()
    return " ".join(words[:10]).strip()


def make_seo_title_options(title: str) -> list[str]:
    base = re.sub(r"\s+", " ", title).strip()
    if not base:
        return []
    opts = [
        trim_at_word_boundary(base, 70),
        trim_at_word_boundary(base + " | Full Report", 70),
        trim_at_word_boundary(base + " | Key Updates", 70),
        trim_at_word_boundary(base + " | News Analysis", 70),
    ]
    out, seen = [], set()
    for x in opts:
        key = x.lower()
        if x and key not in seen:
            out.append(x)
            seen.add(key)
    return out[:4]


def make_meta_options(intro: str, title: str) -> list[str]:
    source = re.sub(r"\s+", " ", intro if intro else title).strip()
    if not source:
        return []
    opts = [
        trim_at_word_boundary(source, 160),
        trim_at_word_boundary(title + " — " + source, 160),
        trim_at_word_boundary("Read the latest details: " + source, 160),
    ]
    out, seen = [], set()
    for x in opts:
        key = x.lower()
        if x and key not in seen:
            out.append(x)
            seen.add(key)
    return out[:3]


def normalize_phrase(text: str, title_case: bool = False) -> str:
    text = re.sub(r"[_-]+", " ", text or "")
    text = re.sub(r"\s+", " ", text).strip(" .:;|-_")
    if not text:
        return ""
    low = text.lower()
    low = re.sub(
        r"\b(?:untitled|design|image|photo|copy|edited|edit|thumb|thumbnail|final|new|jpeg|jpg|png|webp)\b",
        " ",
        low,
    )
    low = re.sub(r"\b\d{2,}\b", " ", low)
    low = re.sub(r"\s+", " ", low).strip()
    if not low:
        low = text.strip()

    if title_case:
        small = {"and", "or", "with", "at", "in", "on", "for", "to", "of", "the", "a", "an"}
        words = []
        for i, word in enumerate(low.split()):
            if i > 0 and word in small:
                words.append(word)
            else:
                words.append(word.capitalize())
        return " ".join(words)
    return low


def trim_words(text: str, max_words: int) -> str:
    words = [w for w in (text or "").split() if w]
    return " ".join(words[:max_words]).strip()


def trim_chars(text: str, max_chars: int) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip(" ,.-:;") + "..."


def infer_subject_from_article(h1: str, focus_keyphrase: str) -> str:
    base = h1 or focus_keyphrase or "news"
    return normalize_phrase(trim_words(base, 8), title_case=False)


def infer_action_from_article(h1: str, intro: str) -> str:
    text = (h1 + " " + intro).lower()
    verbs = [
        "launch", "announces", "announce", "opens", "expands", "wins",
        "releases", "starts", "reveals", "shares", "introduces", "confirms",
        "shows", "updates", "reports"
    ]
    for verb in verbs:
        if verb in text:
            return verb
    return "update"


def infer_context_from_article(intro: str, h1: str, scene_notes: str) -> str:
    if scene_notes.strip():
        return normalize_phrase(scene_notes, title_case=False)
    base = intro or h1 or "news image"
    return normalize_phrase(trim_words(base, 10), title_case=False)


def build_image_alt_text(subject: str, action: str, context: str) -> str:
    text = f"{subject} {action} {context}".strip()
    return trim_chars(re.sub(r"\s+", " ", text), 125)


def build_image_title(subject: str, context: str) -> str:
    text = f"{normalize_phrase(subject, title_case=True)} - {normalize_phrase(context, title_case=True)}"
    return trim_chars(text, 90)


def build_image_caption(subject: str, action: str, context: str) -> str:
    text = f"{normalize_phrase(subject, title_case=True)} {action} {context}".strip()
    return trim_chars(text, 140)


def esc(value: Any) -> str:
    return html.escape(str(value), quote=True)


def image_bytes_to_data_uri(image_bytes: bytes, mime: str = "image/jpeg") -> str:
    encoded = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime};base64,{encoded}"


def compress_image_under_target(img: Image.Image, target_kb: int = 100) -> bytes:
    img = img.convert("RGB")
    best = None

    for width in [img.width, int(img.width * 0.92), int(img.width * 0.85), int(img.width * 0.75), int(img.width * 0.65)]:
        width = max(1, width)
        height = max(1, int(img.height * (width / img.width)))
        resized = img.resize((width, height), Image.LANCZOS)

        for quality in [92, 86, 80, 74, 68, 62, 56, 50, 44]:
            buffer = io.BytesIO()
            resized.save(buffer, format="JPEG", quality=quality, optimize=True)
            data = buffer.getvalue()
            if len(data) <= target_kb * 1024:
                return data
            best = data

    return best if best else b""


def parse_base64_image(data_url: str) -> bytes:
    if "," not in data_url:
        raise ValueError("Invalid image data")
    _, encoded = data_url.split(",", 1)
    return base64.b64decode(encoded)


def build_wordpress_html_fragment(result: dict[str, Any], featured_image_data_uri: str = "") -> str:
    h1 = result.get("h1", "")
    intro = result.get("intro", "")
    structure = result.get("structure", [])
    parts = []

    alt_text = result.get("alt_text", "")
    img_title = result.get("img_title", "")
    caption = result.get("caption", "")

    if featured_image_data_uri:
        img_html = (
            f'<figure class="wp-block-image size-full featured-image-wrap">'
            f'<img src="{featured_image_data_uri}" alt="{esc(alt_text or h1 or "Featured image")}" '
            f'title="{esc(img_title or h1 or "Featured image")}" />'
        )
        if caption:
            img_html += f"<figcaption>{esc(caption)}</figcaption>"
        img_html += "</figure>"
        parts.append(img_html)

    if h1:
        parts.append(f"<h1>{esc(h1)}</h1>")
    if intro:
        parts.append(f"<p>{esc(intro)}</p>")

    for sec in structure:
        if sec.get("h2"):
            parts.append(f"<h2>{esc(sec['h2'])}</h2>")
        for sub in sec.get("subsections", []):
            if sub.get("h3"):
                parts.append(f"<h3>{esc(sub['h3'])}</h3>")
            if sub.get("h4"):
                parts.append(f"<h4>{esc(sub['h4'])}</h4>")
            for p in [x.strip() for x in sub.get("body", "").split("\n\n") if x.strip()]:
                parts.append(f"<p>{esc(p).replace(chr(10), '<br>')}</p>")

    return "\n".join(parts).strip()


def build_full_html_document(result: dict[str, Any], featured_image_data_uri: str = "") -> str:
    h1 = result.get("h1", "")
    intro = result.get("intro", "")
    structure = result.get("structure", [])
    parts = []

    alt_text = result.get("alt_text", "")
    img_title = result.get("img_title", "")
    caption = result.get("caption", "")

    if featured_image_data_uri:
        img_html = (
            f'<figure class="featured-figure"><img src="{featured_image_data_uri}" '
            f'alt="{esc(alt_text or h1 or "Featured image")}" '
            f'title="{esc(img_title or h1 or "Featured image")}" class="featured-image">'
        )
        if caption:
            img_html += f'<figcaption class="featured-caption">{esc(caption)}</figcaption>'
        img_html += "</figure>"
        parts.append(img_html)

    if h1:
        parts.append(f"<h1>{esc(h1)}</h1>")
    if intro:
        parts.append(f'<div class="intro-card"><p class="intro">{esc(intro)}</p></div>')

    for sec in structure:
        section_parts = [f"<h2>{esc(sec['h2'])}</h2>"]
        for sub in sec["subsections"]:
            if sub["h3"]:
                section_parts.append(f"<h3>{esc(sub['h3'])}</h3>")
            if sub["h4"]:
                section_parts.append(f"<h4>{esc(sub['h4'])}</h4>")
            for p in [x.strip() for x in sub["body"].split("\n\n") if x.strip()]:
                section_parts.append(f"<p>{esc(p).replace(chr(10), '<br>')}</p>")
        parts.append('<section class="story-section">' + "".join(section_parts) + "</section>")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(h1 or "SEO Output")}</title>
<style>
:root {{
    --bg: #f4f7fb;
    --card: #ffffff;
    --text: #1f2937;
    --muted: #4b5563;
    --border: #dbe4f0;
    --accent: #2563eb;
}}
* {{ box-sizing: border-box; }}
body {{
    margin: 0;
    font-family: "Segoe UI", Arial, Helvetica, sans-serif;
    color: var(--text);
    background: linear-gradient(180deg, #eef4ff 0%, #f8fbff 100%);
    line-height: 1.8;
    font-size: 18px;
    padding: 28px 18px 48px;
}}
.article-shell {{
    max-width: 940px;
    margin: 0 auto;
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: 22px;
    box-shadow: 0 18px 50px rgba(15, 23, 42, 0.08);
    padding: 28px 28px 34px;
}}
.featured-figure {{ margin: 0 0 24px 0; }}
.featured-image {{
    display: block;
    width: 100%;
    max-width: 100%;
    border-radius: 18px;
}}
.featured-caption {{
    font-size: 14px;
    color: var(--muted);
    margin-top: 8px;
}}
h1 {{
    font-size: 44px;
    line-height: 1.15;
    font-weight: 800;
    color: #0f172a;
    margin: 0 0 18px 0;
}}
.intro-card {{
    background: #f8fbff;
    border: 1px solid var(--border);
    border-left: 5px solid var(--accent);
    border-radius: 16px;
    padding: 16px 18px;
    margin: 0 0 24px 0;
}}
.intro {{
    font-size: 19px;
    margin: 0;
    color: #18212f;
}}
.story-section {{ margin: 28px 0 0 0; }}
h2 {{
    font-size: 31px;
    line-height: 1.22;
    font-weight: 800;
    color: #0f172a;
    margin: 0 0 14px 0;
}}
h3 {{
    font-size: 23px;
    line-height: 1.3;
    font-weight: 700;
    color: #172033;
    margin: 20px 0 10px 0;
}}
h4 {{
    font-size: 18px;
    line-height: 1.35;
    font-weight: 700;
    color: #22304a;
    margin: 16px 0 8px 0;
}}
p {{
    font-size: 18px;
    color: #243041;
    margin: 0 0 16px 0;
}}
</style>
</head>
<body>
<div class="article-shell">
{''.join(parts)}
</div>
</body>
</html>"""


def process_article_payload(article: str, scene_notes: str = "") -> dict[str, Any]:
    lines = clean_lines(article)
    lines = strip_internal_seo_lines(lines)

    if not lines:
        raise ValueError("No valid content found")

    h1 = guess_title(lines)
    intro = build_intro(lines)
    sections = split_body_into_sections(lines, num_sections=3)
    structure = build_nested_article_structure(sections)

    focus_keyphrase = make_focus_keyphrase(h1)
    slug = make_slug(h1)
    short_summary = trim_at_word_boundary(re.sub(r"\s+", " ", intro if intro else h1).strip(), 200)

    seo_titles = make_seo_title_options(h1)
    meta_options = make_meta_options(intro, h1)
    seo_title = seo_titles[0] if seo_titles else h1
    meta_description = meta_options[0] if meta_options else trim_at_word_boundary(intro if intro else h1, 160)

    plain_parts = []
    headings_summary = []
    body_blocks = []

    if h1:
        plain_parts += [h1, ""]
    if intro:
        plain_parts += [intro, ""]

    for sec in structure:
        plain_parts += [sec["h2"], ""]
        headings_summary.append(sec["h2"])
        body_blocks.append(sec["h2"])

        sub_bodies = []
        for sub in sec["subsections"]:
            if sub["h3"]:
                headings_summary.append("  - " + sub["h3"])
            if sub["h4"]:
                headings_summary.append("    * " + sub["h4"])
            if sub["body"].strip():
                sub_bodies.append(sub["body"].strip())

            if sub["h3"]:
                plain_parts.append(sub["h3"])
            if sub["h4"]:
                plain_parts.append(sub["h4"])
            plain_parts += [sub["body"], ""]

        if sub_bodies:
            body_blocks.append("\n\n".join(sub_bodies))

    subject = infer_subject_from_article(h1, focus_keyphrase)
    action = infer_action_from_article(h1, intro)
    context = infer_context_from_article(intro, h1, scene_notes)

    alt_text = build_image_alt_text(subject, action, context)
    img_title = build_image_title(subject, context)
    caption = build_image_caption(subject, action, context)

    return {
        "h1": h1,
        "intro": intro,
        "structure": structure,
        "generated_plain_text": "\n".join(plain_parts).strip(),
        "headings_copy": "\n".join(headings_summary),
        "body_copy": "\n\n".join(body_blocks),
        "focus_keyphrase": focus_keyphrase,
        "seo_title": seo_title,
        "meta_description": meta_description,
        "slug": slug,
        "short_summary": short_summary,
        "seo_titles": seo_titles,
        "meta_options": meta_options,
        "alt_text": alt_text,
        "img_title": img_title,
        "caption": caption,
    }


# -----------------------------
# Routes
# -----------------------------
@app.get("/")
def index():
    return render_template("index.html")


@app.post("/process")
def process_article():
    data = request.get_json(force=True)
    article = (data.get("article") or "").strip()
    scene_notes = (data.get("scene_notes") or "").strip()

    if not article:
        return jsonify({"error": "Please paste an article first."}), 400

    try:
        result = process_article_payload(article, scene_notes)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.post("/compress-image")
def compress_image():
    data = request.get_json(force=True)
    image_data = data.get("image_data", "")

    if not image_data:
        return jsonify({"error": "No image received."}), 400

    try:
        raw = parse_base64_image(image_data)
        image = Image.open(io.BytesIO(raw)).convert("RGB")
        compressed = compress_image_under_target(image, target_kb=100)
        if not compressed:
            return jsonify({"error": "Unable to compress image."}), 400

        data_uri = image_bytes_to_data_uri(compressed, "image/jpeg")
        size_kb = round(len(compressed) / 1024, 1)
        return jsonify({
            "image_data": data_uri,
            "size_kb": size_kb
        })
    except Exception as exc:
        return jsonify({"error": f"Image processing failed: {exc}"}), 400


@app.post("/export/txt")
def export_txt():
    data = request.get_json(force=True)
    content = (data.get("content") or "").strip()
    filename = (data.get("filename") or "seo-output").strip()

    if not content:
        return jsonify({"error": "Nothing to export."}), 400

    buffer = io.BytesIO(content.encode("utf-8"))
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{filename}.txt",
        mimetype="text/plain"
    )


@app.post("/export/html")
def export_html():
    data = request.get_json(force=True)
    result = data.get("result") or {}
    image_data = data.get("image_data", "")
    filename = (data.get("filename") or "seo-output").strip()

    if not result:
        return jsonify({"error": "Nothing to export."}), 400

    html_doc = build_full_html_document(result, image_data)
    buffer = io.BytesIO(html_doc.encode("utf-8"))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{filename}.html",
        mimetype="text/html"
    )


@app.post("/export/docx")
def export_docx():
    if not HAS_DOCX:
        return jsonify({"error": "python-docx is not installed on the server."}), 400

    data = request.get_json(force=True)
    result = data.get("result") or {}
    image_data = data.get("image_data", "")
    filename = (data.get("filename") or "seo-output").strip()

    if not result:
        return jsonify({"error": "Nothing to export."}), 400

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    temp_img_path = None

    try:
        if image_data:
            raw = parse_base64_image(image_data)
            image = Image.open(io.BytesIO(raw)).convert("RGB")
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                temp_img_path = tmp.name
            image.save(temp_img_path, format="JPEG", quality=95)
            doc.add_picture(temp_img_path, width=Inches(6.8))
            if result.get("caption"):
                cap = doc.add_paragraph()
                cap.paragraph_format.space_before = Pt(4)
                cap.paragraph_format.space_after = Pt(10)
                run = cap.add_run(result["caption"])
                run.font.name = "Segoe UI"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)

        def add_heading(text: str, size: int, space_before: int = 8, space_after: int = 4) -> None:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            r = p.add_run(text)
            r.bold = True
            r.font.name = "Segoe UI"
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(17, 17, 17)

        def add_para(text: str, size: int = 12) -> None:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(text)
            r.font.name = "Segoe UI"
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(34, 34, 34)

        if result.get("h1"):
            add_heading(result["h1"], 26, 0, 8)
        if result.get("intro"):
            add_para(result["intro"])

        for sec_data in result.get("structure", []):
            add_heading(sec_data["h2"], 22, 12, 6)
            for sub in sec_data.get("subsections", []):
                if sub.get("h3"):
                    add_heading(sub["h3"], 18, 8, 4)
                if sub.get("h4"):
                    add_heading(sub["h4"], 14, 6, 2)
                for p in [x.strip() for x in sub.get("body", "").split("\n\n") if x.strip()]:
                    add_para(p)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name=f"{filename}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    finally:
        if temp_img_path and os.path.exists(temp_img_path):
            try:
                os.remove(temp_img_path)
            except Exception:
                pass


@app.post("/wp-html")
def wp_html():
    data = request.get_json(force=True)
    result = data.get("result") or {}
    image_data = data.get("image_data", "")

    if not result:
        return jsonify({"error": "Generate SEO output first."}), 400

    fragment = build_wordpress_html_fragment(result, image_data)
    return jsonify({"html": fragment})


if __name__ == "__main__":
    app.run(debug=True)
